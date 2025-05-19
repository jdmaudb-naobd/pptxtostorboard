"""
Storyboard Generator
Converts PowerPoint presentations to structured storyboard documents
"""

import os
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Inches
from pptx import Presentation
import json
from sentence_transformers import SentenceTransformer
import numpy as np

from models.image_classifier import HybridImageClassifier
from models.slide_classifier import SlideClassifier
from utils.abbreviation_handler import AbbreviationHandler
from utils.content_validator import ContentValidator
from pptx_extractor import PPTXExtractor

class StoryboardGenerator:
    def __init__(self, config_path: Optional[str] = None):
        """
        Initialize the storyboard generator
        
        Args:
            config_path: Path to configuration file
        """
        # Initialize components
        self.image_classifier = HybridImageClassifier()
        self.slide_classifier = SlideClassifier()
        self.abbreviation_handler = AbbreviationHandler()
        self.content_validator = ContentValidator()
        self.pptx_extractor = None
        
        # Load BGE-M3 for semantic matching
        self.model = SentenceTransformer("BAAI/bge-m3")
        
        # Load configuration if available
        self.config = {
            "template_path": None,
            "instruction_path": None,
            "output_path": "output",
            "training_pairs_path": None
        }
        
        if config_path and os.path.exists(config_path):
            with open(config_path, 'r') as f:
                self.config.update(json.load(f))

    def load_training_data(self):
        """Load and process training examples"""
        if not self.config["training_pairs_path"]:
            return
        
        # TODO: Implement training data processing
        pass

    def process_pptx(self, pptx_path: str) -> Dict:
        """
        Process PowerPoint file and extract structured content
        """
        self.pptx_extractor = PPTXExtractor(pptx_path)
        
        # Extract text and images
        text_content = self.pptx_extractor.extract_text()
        image_info = self.pptx_extractor.extract_images("temp_images")
        
        # Process each slide
        processed_content = []
        for slide_idx, slide in enumerate(text_content):
            # Classify slide
            slide_type = self.slide_classifier.get_slide_type(slide)
            
            # Process text for abbreviations
            highlighted_text, abbreviations = self.abbreviation_handler.highlight_abbreviations(
                slide["text"]
            )
            
            # Validate content
            validation_results = self.content_validator.validate_content(slide["text"])
            
            # Get images for this slide
            slide_images = [
                img for img in image_info
                if img["slide_number"] == slide_idx + 1
            ]
            
            # Classify images
            for img in slide_images:
                img["semantic_type"] = self.image_classifier.get_image_metadata(img["path"])
            
            processed_content.append({
                "slide_number": slide["slide_number"],
                "slide_type": slide_type,
                "text": highlighted_text,
                "abbreviations": abbreviations,
                "images": slide_images,
                "validation_results": validation_results
            })
        
        return processed_content

    def generate_storyboard(self, processed_content: Dict, output_path: str):
        """
        Generate storyboard document from processed content
        """
        # Create new document or load template
        doc = Document(self.config["template_path"]) if self.config["template_path"] else Document()
        
        # Create contents chapters table
        self._create_contents_table(doc, processed_content)
        
        # Create abbreviations table
        self._create_abbreviations_table(doc)
        
        # Create content tables for each slide
        for slide in processed_content:
            self._create_slide_content_table(doc, slide)
        
        # Create question tables if applicable
        for slide in processed_content:
            if slide["slide_type"][0] == "quiz":
                self._create_question_table(doc, slide)
        
        # Save document
        doc.save(output_path)

    def _create_contents_table(self, doc: Document, content: Dict):
        """Create table of contents"""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "Chapter"
        table.rows[0].cells[1].text = "Subchapter"
        
        # TODO: Implement chapter organization logic

    def _create_abbreviations_table(self, doc: Document):
        """Create abbreviations table"""
        abbrev_list = self.abbreviation_handler.create_abbreviations_table()
        if abbrev_list:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = "Abbreviation"
            table.rows[0].cells[1].text = "Definition"
            
            for abbrev in abbrev_list:
                row = table.add_row()
                row.cells[0].text = abbrev["abbreviation"]
                row.cells[1].text = abbrev["definition"]

    def _create_slide_content_table(self, doc: Document, slide: Dict):
        """Create content table for a slide"""
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        # Fill in table cells
        cells = [
            ("Chapter", ""),  # To be filled based on content organization
            ("Subchapter", ""),
            ("Text", slide["text"]),
            ("Media/Images", self._format_image_info(slide["images"])),
            ("Visual Details", self._format_visual_details(slide)),
            ("Interactivity Details", ""),  # To be filled based on slide type
            ("References", ""),  # To be extracted from slide content
            ("Extra Details/Settings", "")
        ]
        
        for idx, (header, content) in enumerate(cells):
            table.rows[idx].cells[0].text = header
            table.rows[idx].cells[1].text = content

    def _create_question_table(self, doc: Document, slide: Dict):
        """Create table for quiz questions"""
        if slide["slide_type"][0] != "quiz":
            return
            
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        # TODO: Implement question table formatting

    def _format_image_info(self, images: List[Dict]) -> str:
        """Format image information for the table"""
        if not images:
            return "No images"
            
        return "\n".join(
            f"Image {idx+1}: {img['semantic_type']['predicted_type']} "
            f"({img['dimensions']})"
            for idx, img in enumerate(images)
        )

    def _format_visual_details(self, slide: Dict) -> str:
        """Format visual details for the table"""
        details = []
        if slide["images"]:
            details.append("Contains visual elements:")
            for img in slide["images"]:
                details.append(f"- {img['semantic_type']['predicted_type']}")
        
        return "\n".join(details) if details else "No specific visual details"

    def save_config(self, path: str):
        """Save current configuration"""
        with open(path, 'w') as f:
            json.dump(self.config, f, indent=2)

    def load_config(self, path: str):
        """Load configuration"""
        with open(path, 'r') as f:
            self.config.update(json.load(f)) 